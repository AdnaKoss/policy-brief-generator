"""Renders a PolicyBrief as a .docx file in a plain UN/UNDP-style layout:
a title, a dated byline, and headed sections for each part of the brief."""

from __future__ import annotations

import io

from docx import Document

from generation.schemas import PolicyBrief


def _add_paragraphs(document: Document, text: str) -> None:
    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            document.add_paragraph(block)


def _add_list(document: Document, items: list[str], style: str) -> None:
    for item in items:
        document.add_paragraph(item, style=style)


def build_docx(brief: PolicyBrief) -> bytes:
    document = Document()

    document.add_heading(brief.title, level=0)
    document.add_paragraph().add_run(brief.date).italic = True

    document.add_heading("Executive Summary", level=1)
    _add_paragraphs(document, brief.executive_summary)

    document.add_heading("Background", level=1)
    _add_paragraphs(document, brief.background)

    document.add_heading("Key Findings", level=1)
    _add_list(document, brief.key_findings, style="List Bullet")

    document.add_heading("Policy Implications", level=1)
    _add_paragraphs(document, brief.policy_implications)

    document.add_heading("Recommendations", level=1)
    _add_list(document, brief.recommendations, style="List Number")

    document.add_heading("Sources", level=1)
    _add_list(document, brief.sources, style="List Bullet")

    if brief.limitations:
        document.add_heading("Limitations", level=1)
        _add_list(document, brief.limitations, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
