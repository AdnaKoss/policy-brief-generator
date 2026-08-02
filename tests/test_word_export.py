from __future__ import annotations

import io

from docx import Document

from generation.schemas import PolicyBrief
from templates.word_export import build_docx


def _make_brief(limitations: list[str] | None = None) -> PolicyBrief:
    return PolicyBrief(
        title="AI Governance in Practice",
        executive_summary="Summary paragraph.",
        background="First background paragraph.\n\nSecond background paragraph.",
        key_findings=["Finding one.", "Finding two."],
        policy_implications="Implications paragraph.",
        recommendations=["Establish a taskforce.", "Require impact assessments."],
        sources=["https://example.org/report"],
        limitations=limitations if limitations is not None else [],
        date="2026-08-03",
    )


def _paragraph_texts(document: Document) -> list[str]:
    return [p.text for p in document.paragraphs]


def test_build_docx_includes_title_and_date():
    brief = _make_brief()
    document = Document(io.BytesIO(build_docx(brief)))

    assert document.paragraphs[0].text == "AI Governance in Practice"
    assert document.paragraphs[0].style.name == "Title"
    assert "2026-08-03" in _paragraph_texts(document)


def test_build_docx_includes_all_section_headings():
    brief = _make_brief()
    document = Document(io.BytesIO(build_docx(brief)))

    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    assert headings == [
        "Executive Summary",
        "Background",
        "Key Findings",
        "Policy Implications",
        "Recommendations",
        "Sources",
    ]


def test_build_docx_splits_multi_paragraph_background():
    brief = _make_brief()
    document = Document(io.BytesIO(build_docx(brief)))
    texts = _paragraph_texts(document)

    assert "First background paragraph." in texts
    assert "Second background paragraph." in texts


def test_build_docx_uses_bullets_and_numbers_correctly():
    brief = _make_brief()
    document = Document(io.BytesIO(build_docx(brief)))

    findings = [p for p in document.paragraphs if p.text in brief.key_findings]
    assert all(p.style.name == "List Bullet" for p in findings)

    recommendations = [p for p in document.paragraphs if p.text in brief.recommendations]
    assert all(p.style.name == "List Number" for p in recommendations)


def test_build_docx_omits_limitations_heading_when_empty():
    brief = _make_brief(limitations=[])
    document = Document(io.BytesIO(build_docx(brief)))

    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    assert "Limitations" not in headings


def test_build_docx_includes_limitations_when_present():
    brief = _make_brief(limitations=["Further research needed on X."])
    document = Document(io.BytesIO(build_docx(brief)))

    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    texts = _paragraph_texts(document)
    assert "Limitations" in headings
    assert "Further research needed on X." in texts
